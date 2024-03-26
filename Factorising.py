from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import os

TITLE = "Factorising Quadractics"

def FactorisingQuadratics(numQ):
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),'3')
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE+" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')
    count = 0
    while count != numQ:

        a = choice([i for i in range(-9,9) if i not in [0]])
        
        b = choice([i for i in range(-9,9) if i not in [0]])
        
        c = a + b
        if c < 0:
            op1 = "-"
            c = c*(-1)
        else:
            op1 = "+"
        if c == 1:
            c = ""
        d = a*b
        if d < 0:
            op2 = "-"
            d = d*(-1)
        else:
            op2 = "+"
        if c != 0:
            question = " " + op1 + " " + str(c) + "x " + op2 + " " + str(d)
        else:
            question = " " + op2 + " " + str(d)
        if a < 0:
            ansop1 = "-"
            a = a*(-1)
        else:
            ansop1 = "+"
        if b < 0:
            ansop2 = "-"
            b = b*(-1)
        else:
            ansop2 = "+"
        
        ans = "(x " + ansop1 + str(a) + ")(x " + ansop2 + str(b) + ")"
        
        if count%8 == 0 and count != 0:
            documentQ.add_paragraph()
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        pq = documentQ.add_paragraph("x")
        subscript = pq.add_run("2")
        subscript.font.superscript = True
        subscript.font.size = Pt(12)
        pq.add_run(question)
        documentQ.add_paragraph()

        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans))
        #subscript = p.add_run("")
        #subcript.font.superscript = True
        count += 1
    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
        
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)




    
