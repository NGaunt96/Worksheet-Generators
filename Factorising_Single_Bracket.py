from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import os

TITLE = "Factorising Single Brackets"
def flipSign(sign, pos):
    if pos == "in":
        if sign == "-":
            return "+"
        else:
            return "-"
    else:
        if sign == "-":
            return ""
        else:
            return "-"
def FactorisingSingleBracket(numQ):
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),'3')
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE+" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')
    count = 0
    while count != numQ:
        factor = choice([i for i in range(1, 10)])
        if choice([1,2,3,4,5]) in [1, 2]:
            factor = factor*-1
        a = choice([i for i in range(1, 7)])
        if choice([1,2,3,4,5]) in [1, 2]:
            a = a*-1
        b = choice([i for i in range(1, 7)])
        if choice([1,2,3,4,5]) in [1, 2]:
            b = b*-1
        #x coefficient and constant
        c = factor*a
        d = factor*b

        if factor < 0:
            factor = factor*-1
            ansop1 = "-"
        else:
            ansop1 = ""
        if factor == 1:
            factor = ""

        if b < 0:
            ansop3 = "-"
            b = b *-1
        else:
            ansop3 = "+"
        if a < 0:
            ansop2 = ""
            a = a *-1
            ansop1 = flipSign(ansop1, "out")
            ansop3 = flipSign(ansop3, "in")
        else:
            ansop2 = ""
        if a == 1:
            a = ""
        if c < 0:
            quop1 = "-"
            c = c * -1
        else:
            quop1 = ""
        if c == 1:
            c = ""
        if d < 0:
            quop2 = "-"
            d = d*-1
        else:
            quop2 = "+"
        ans = ansop1 + str(factor) + "(" + ansop2 + str(a) + "x" + ansop3 + str(b) + ")"
        question = quop1 + str(c) + "x" + quop2 + str(d)
        
        if count%8 == 0 and count != 0:
            documentQ.add_paragraph()
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        pq = documentQ.add_paragraph(question)
        documentQ.add_paragraph()

        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans))
        #subscript = p.add_run("")
        #subcript.font.superscript = True
        count += 1
    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
        
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)

FactorisingSingleBracket(1000)


    
