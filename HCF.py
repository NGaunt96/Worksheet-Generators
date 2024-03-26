from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

TITLE = "HCF"

#initialise primes list. 2's have a higher weight than 3 and 5 which has a higher weight than 7
primes = [2, 3, 5, 7]
weightedprimes = [2, 2, 2, 2, 2, 3, 3, 3, 5, 5, 5, 7]

def HCF(numQ):
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),'2')
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE + " Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')
    count = 0
    while count != numQ:
        chosenprimesa = []
        chosenprimesb = []
        hcfprimes = []
        lcmprimes = []
        questions = ["HCF", "LCM"]
        #q = questions[randint(0,len(questions)-1)]
        q = "HCF"
        #Choose some primes for number a
        for i in range(randint(3, 5)):
           chosenprimesa.append(weightedprimes[randint(0,len(weightedprimes)-1)])
        #select a random assortment of them to be part of number b.
        for i in range(randint(1, len(chosenprimesa)-2)):
           sharer = chosenprimesa[randint(0,len(chosenprimesa)-1)]
           chosenprimesb.append(sharer)
        #create a difference in how many primes and fill in the primes of b
        diff = randint(0, len(chosenprimesa))
        while len(chosenprimesb) < diff:
           chosenprimesb.append(weightedprimes[randint(0,len(weightedprimes)-1)])

        #work out value of a and b
        a = 1
        b = 1
        for i in range(0, len(chosenprimesa)):
           a = a * chosenprimesa[i]
        for i in range(0, len(chosenprimesb)):
           b = b * chosenprimesb[i]
        if b in primes:
            fixer = weightedprimes[randint(0, len(weightedprimes)-1)]
            b = b * fixer
            chosenprimesb.append(fixer)
            
        #find out which primes are in both primes' lists.
        counta = {}
        countb = {}
        for i in primes:
           counta[i] = 0
           countb[i] = 0

        for i in chosenprimesa:
           counta[i] += 1
        for i in chosenprimesb:
           countb[i] += 1

        for i in primes:
           while counta[i] > 0 and countb[i] > 0:
              hcfprimes.append(i)
              counta[i] -= 1
              countb[i] -= 1

        #work out HCF
        hcf = 1
        for i in hcfprimes:
           hcf = hcf*i

        #use remaining primeslists to find primes for LCM
        for i in primes:
           while counta[i] > 0:
              lcmprimes.append(i)
              counta[i] -= 1
           while countb[i] > 0:
              lcmprimes.append(i)
              countb[i] -= 1
        #work out LCM
        lcm = 1
        for i in lcmprimes:
           lcm = lcm*i
        lcm = lcm * hcf

        if q == "HCF":
            ans = hcf
        else:
            ans = lcm

        
        question = "What is the " + q +  " of " + str(a) + " and " + str(b) + "?"
        if count%8 == 0 and count != 0:
            documentQ.add_paragraph()
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        documentQ.add_paragraph(question)
        documentQ.add_paragraph()

        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans))
        if q == "area":
            squared = p.add_run("2")
            squared.font.superscript = True
        count += 1

    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
    
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)




    
