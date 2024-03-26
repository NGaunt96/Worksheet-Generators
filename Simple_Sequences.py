from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
#---This is the title of the worksheet and also the name of the question and answer sheets---#
TITLE = "Finishing Sequences"


questionsColumns = '2'

def Simple_Sequences(numQ):
    
    #---Creates two documents one for the questions, one for the answers questions columns are how many columns the worksheet is split into---#
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),questionsColumns)
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE+" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')

    
    count = 0
    while count != numQ:
        if count < numQ/5:
            add = randint(1,6)
        else:
            add = randint(7,13)
        start = randint(1, 15)
        seq = [start]
        for i in range(10):
            seq.append(seq[-1] + add)
        ans = ""
        for i in seq:
            if i == seq[-1]:
                ans = ans + str(i)
            else:
                ans = ans + str(i) + ", "
        for i in range(choice([4,5,6,7,8])):
            pos = randint(0,len(seq)-1)
            while seq[pos] == "_":
                pos = randint(0,len(seq)-1)
            seq[pos] = "_"
        question = ""
        for i in range(len(seq)-1):
            if i == len(seq)-2:
                question = question + str(seq[i])
            else:
                question = question + str(seq[i]) + ", "
        #---If the questions are all the same number of lines, format it to go to the next page after n questions---#
        if count%8 == 0 and count != 0:
            documentQ.add_paragraph()
        #---Add the generated question and answer into the respective worksheet---#
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        documentQ.add_paragraph(question)
        documentQ.add_paragraph()

        if count%8 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans))
        #---These two lines can be used to add supercript for squared or cubed numbers or units.---#
        #subscript = p.add_run("2")
        #subcript.font.superscript = True
        count += 1
        
    #---Save the two worksheets---#
    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
    
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)





    
