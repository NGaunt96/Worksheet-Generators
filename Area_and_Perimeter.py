from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

TITLE = "Area and Perimeter"


def AreaAndPerimeter(numQ):
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),'2')
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE +" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')

    
    count = 0
    while count != numQ:
        shapes = ["right-angled triangle", "rectangle", "square", "parallelogram"]
        questions = ["area", "perimeter"]
        #Pick the question. eg area of square
        shape = shapes[randint(0,len(shapes)-1)]
        q = questions[randint(0,len(questions)-1)]

        #depending on the shape, assign values, work out answers, and create question
        if shape == "right-angled triangle":
            height = randint(1, 20)
            width = randint(1, 20)
            hyp = round(sqrt(height**2+width**2))
            if q == "area":
                ans = 0.5*height*width
            if q == "perimeter":
                ans = height + width + hyp
            question = "What is the " + q + " of a " + shape + " with height of " + str(height) + "cm, a width of " + str(width) + "cm and a hypotenuse of " + str(hyp) + "cm?"
            
        elif shape == "rectangle":
            length = randint(1, 20)
            width = randint(1, 20)
            while width == length:
                width = randint(1, 20)
            if q == "area":
                ans = length*width
            if q == "perimeter":
                ans = (length + width)*2
            question = "What is the " + q + " of a " + shape + " with length of " + str(length) + "cm and a width of " + str(width) + "cm?"
            

        
        elif shape == "square":
            side = randint(1, 20)
            if q == "area":
                ans = side**2
            if q == "perimeter":
                ans = side*4
            question = "What is the " + q + " of a " + shape + " with side length of " + str(side) + "cm?"
        
        elif shape == "parallelogram":
            length = randint(1, 20)
            width = randint(1, 20)
            while width == length:
                width = randint(1, 20)
            if q == "area":
                ans = length*width
            if q == "perimeter":
                ans = (length + width)*2
            question = "What is the " + q + " of a " + shape + " with length of " + str(length) + "cm and a width of " + str(width) + "cm?"

    ##    if shape == "trapezium":
    ##        q = "area"
    ##        a = randint(1, 20)
    ##        b = randint(1, 20)
    ##        while b == a:
    ##            b = randint(1, 20)

        #if count%12 == 0 and count != 0:
            #documentQ.add_paragraph()
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        documentQ.add_paragraph(question)
        documentQ.add_paragraph()

        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans) +"cm")
        if q == "area":
            squared = p.add_run("2")
            squared.font.superscript = True
        count += 1
    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
    
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)




    
