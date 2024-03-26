from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
#---This is the title of the worksheet and also the name of the question and answer sheets---#
TITLE = "Expanding Cubics"


questionsColumns = '2'

def ExpandingCubics(numQ):

    #---Creates two documents one for the questions, one for the answers questions columns are how many columns the worksheet is split into---#
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),questionsColumns)
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE+" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')

    
    count = 0
    while count != numQ:
        a = choice([1, 1, 2])
        c = choice([1, 1, 2])
        e = choice([1, 1, 2])
        b = choice([i for i in range(1,5)])
        d = choice([i for i in range(1,5)])
        f = choice([i for i in range(1,5)])
        while b == d or b == f:
            b = choice([i for i in range(1,5)])
        while d == f or d == b:
            d = choice([i for i in range(1,5)])
        coeff1 = a*c*e
        coeff2 = a*c*f + a*d*e + b*c*e
        coeff3 = a*d*f + b*c*f + b*d*e
        const = b*d*f
        
        if coeff1 == 1:
            coeff1 == ""
        if coeff2 == 1:
            coeff2 == ""
        if coeff3 == 1:
            coeff3 == ""
        if a == 1:
            a = ""
        if c == 1:
            c = ""
        if e == 1:
            e = ""
        if coeff1 == -1:
            coeff1 == "-"
        if coeff2 == -1:
            coeff2 == "-"
        if coeff3 == -1:
            coeff3 == "-"
        if a == -1:
            a = "-"
        if c == -1:
            c = "-"
        if e == -1:
            e = "-"
        ans1 = str(coeff1) + "x"
        ans2 = " + " + str(coeff2) + "x"
        ans3 = " + " + str(coeff3) + "x + " + str(const)
        question = "(" + str(a) + "x + " + str(b) + ")" + "(" + str(c) + "x + " + str(d) + ")" +"(" + str(e) + "x + " + str(f) + ")" 

        #---If the questions are all the same number of lines, format it to go to the next page after n questions---#
        if count%8 == 0 and count != 0:
            documentQ.add_paragraph()
        #---Add the generated question and answer into the respective worksheet---#
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        documentQ.add_paragraph(question)
        documentQ.add_paragraph()

        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans1))
        #---These two lines can be used to add supercript for squared or cubed numbers or units.---#
        subscript1 = p.add_run("3")
        subscript1.font.superscript = True
        answerPart2 = p.add_run(ans2)
        subscript2 = p.add_run("2")
        subscript2.font.superscript = True
        answerPart2 = p.add_run(ans3)
        count += 1
        
    #---Save the two worksheets---#
    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
    
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)



ExpandingCubics(1000)
    
