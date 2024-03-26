from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

TITLE = "Red Percentage Increase and Decrease"
questions = ["Increase", "Decrease"]

def RandPercInDe(numQ):
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),'3')
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE+" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')
    count = 0
    while count != numQ:
        q = choice(questions)
        perc = randint(1, 100)
        n = randint(10, 300)
        if choice([1, 2]) == 1:
            n = n*10
        question = q + " " + str(n) + " by " + str(perc) + "%"
        if q == "Increase":
            ans = n + n*(perc/100)
        else:
            ans = n - n*(perc/100)
        
        #if count%12 == 0 and count != 0:
            #documentQ.add_paragraph()
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        documentQ.add_paragraph(question)
        documentQ.add_paragraph()

        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans))
        #subscript = p.add_run("2")
        #subcript.font.superscript = True
        count += 1

    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
    
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)





    
