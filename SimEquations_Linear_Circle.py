from random import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

TITLE = "Solving Simultaneous Equations - 1 Linear, 1 Circle"

documentQ = Document()
sectionQ = documentQ.sections[0]
headerQ = sectionQ.header
paragraphQ = headerQ.paragraphs[0]
paragraphQ.text = TITLE
paragraphQ.style = documentQ.styles["Title"]
sectPrQ = sectionQ._sectPr
colsQ = sectPrQ.xpath('./w:cols')[0]
colsQ.set(qn('w:num'),'3')

documentAns = Document()
sectionAns = documentAns.sections[0]
headerAns = sectionAns.header
paragraphAns = headerAns.paragraphs[0]
paragraphAns.text = TITLE + " Answers"
paragraphAns.style = documentAns.styles["Title"]
sectPrAns = sectionAns._sectPr
colsAns = sectPrAns.xpath('./w:cols')[0]
colsAns.set(qn('w:num'),'3')

def SolveLinearSimultaneousLinCir(numQ):
  count = 0
  while count != numQ:
    y = choice([i for i in range(1,5)])
    a = choice([i for i in range(1,4)])
    x = choice([i for i in range(1,5)])

    if choice([1, 2, 3, 4, 5]) == 1:
        a = a *(-1)
    if choice([1, 2, 3, 4, 5]) == 1:
        x = x *(-1)
    r = y**2+x**2
    b = y-a*x
    if a == 1:
        a = ""
    if a == -1:
        a = "-"
    if b < 0:
        op1 = "-"
        b = b*(-1)
    else:
        op1 = "+"
    eq1p1 = "x"
    eq1p2 = " + y"
    eq1p3 = " = " + str(r)
    eq2 = "y = " + str(a) + "x " + op1 + " " + str(b)
    #if count%8 == 0 and count != 0:
    #  documentQ.add_paragraph()
    #if count%12 == 0 and count != 0:
    #  documentAns.add_paragraph()
    paraQ = documentQ.add_paragraph()
    runQ = paraQ.add_run()
    runQ.text = "Question " + str(count+1)
    runQ.underline = True
    p = documentQ.add_paragraph(eq1p1)
    subscript1 = p.add_run("2")
    subscript1.font.superscript = True
    questionPart2 = p.add_run(eq1p2)
    subscript2 = p.add_run("2")
    subscript2.font.superscript = True
    questionPart2 = p.add_run(eq1p3)
    documentQ.add_paragraph(eq2)

    paraAns = documentAns.add_paragraph()
    runAns = paraAns.add_run()
    runAns.text = "Question " + str(count+1)
    runAns.underline = True
    documentAns.add_paragraph("x = " + str(x) + " and y = " + str(y))
    count += 1
    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")

    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)
     
SolveLinearSimultaneousLinCir(10)
