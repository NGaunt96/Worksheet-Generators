from random import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

TITLE = "Solving Linear Simultaneous Equations"

documentQ = Document()
sectionQ = documentQ.sections[0]
headerQ = sectionQ.header
paragraphQ = headerQ.paragraphs[0]
paragraphQ.text = TITLE
paragraphQ.style = documentQ.styles["Title"]
sectPrQ = sectionQ._sectPr
colsQ = sectPrQ.xpath('./w:cols')[0]
colsQ.set(qn('w:num'),'3')

documentAns = Document()
sectionAns = documentAns.sections[0]
headerAns = sectionAns.header
paragraphAns = headerAns.paragraphs[0]
paragraphAns.text = TITLE + " Answers"
paragraphAns.style = documentAns.styles["Title"]
sectPrAns = sectionAns._sectPr
colsAns = sectPrAns.xpath('./w:cols')[0]
colsAns.set(qn('w:num'),'3')

def SolveLinearSimultaneous(numQ):
  count = 0
  while count != numQ:
     x = choice([i for i in range(-10,10) if i not in [0]])
     y = choice([i for i in range(-10,10) if i not in [0]])
     while x == y:
        y = randint(1,9)
     coeff1 = randint(1,5)
     coeff2 = randint(1,5)
     coeff3 = randint(1,5)
     coeff4 = randint(1,5)
     op1 = randint(1,2)
     op2 = randint(1,2)
     
     if op1 == 1:
        op1 = "+"
     else:
        op1 = "-"

     if op2 == 1:
        op2 = "+"
     else:
        op2 = "-"

     eq1 = "a = " + str(coeff1) + "*" + str(x) + op1 + str(coeff2) + "*" + str(y)
     eq2 = "b = " + str(coeff3) + "*" + str(x) + op2 + str(coeff4) + "*" + str(y)
     ans1 = exec(eq1)
     ans2 = exec(eq2)
     if coeff1 == 1:
          coeff1 = ""
     if coeff2 == 1:
          coeff2 = ""
     if coeff3 == 1:
          coeff3 = ""
     if coeff4 == 1:
          coeff4 = ""
     eq1 = "(1)  " + str(coeff1) + "x " + op1 + " " + str(coeff2) + "y = " + str(a)
     eq2 = "(2)  " + str(coeff3) + "x " + op2 + " " + str(coeff4) + "y = " + str(b)

     if count%8 == 0 and count != 0:
        documentQ.add_paragraph()
     if count%12 == 0 and count != 0:
        documentAns.add_paragraph()
     paraQ = documentQ.add_paragraph()
     runQ = paraQ.add_run()
     runQ.text = "Question " + str(count+1)
     runQ.underline = True
     documentQ.add_paragraph(eq1)
     documentQ.add_paragraph(eq2)

     paraAns = documentAns.add_paragraph()
     runAns = paraAns.add_run()
     runAns.text = "Question " + str(count+1)
     runAns.underline = True
     documentAns.add_paragraph("x = " + str(x) + " and y = " + str(y))
     count += 1
     directory = os.getcwd()
     os.chdir(directory+'/'+"Worksheets")

     documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
     documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
     os.chdir(directory)
     
