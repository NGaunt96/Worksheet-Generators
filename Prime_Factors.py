from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

TITLE = "Prime Factors"

#initialise primes list. 2's have a higher weight than 3 and 5 which has a higher weight than 7
primes = [2, 3, 5, 7]
weightedprimes = [2, 2, 2, 2, 2, 3, 3, 3, 5, 5, 5, 7]


def PrimeFactors(numQ):
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),'2')
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE + " Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')

    count = 0
    while count != numQ:
        chosenprimes = []
        for i in range(randint(3,6)):
            chosenprimes.append(choice(weightedprimes))
        #work out value of the answer
        a = 1
        ans = ""
        tempcount = 0
        for i in chosenprimes:
            a *= i
            ans += str(i)
            if tempcount != len(chosenprimes)-1:
                ans += " x "
            tempcount += 1
        
        question = "What is the prime factorisation of " + str(a) + "?"
        if count%8 == 0 and count != 0:
            documentQ.add_paragraph()
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        documentQ.add_paragraph(question)
        documentQ.add_paragraph()

        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans))
        count += 1

    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
    
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)




    
