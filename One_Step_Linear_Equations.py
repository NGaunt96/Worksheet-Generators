from random import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

TITLE = "Solving One Step Linear Equations"



letters = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z']
ops = ["+", "-", "÷", "×"]
def OneStepLinearEquations(numQ):
  documentQ = Document()
  sectionQ = documentQ.sections[0]
  headerQ = sectionQ.header
  paragraphQ = headerQ.paragraphs[0]
  paragraphQ.text = TITLE
  paragraphQ.style = documentQ.styles["Title"]
  sectPrQ = sectionQ._sectPr
  colsQ = sectPrQ.xpath('./w:cols')[0]
  colsQ.set(qn('w:num'),'3')
     
  documentAns = Document()
  sectionAns = documentAns.sections[0]
  headerAns = sectionAns.header
  paragraphAns = headerAns.paragraphs[0]
  paragraphAns.text = TITLE + " Answers"
  paragraphAns.style = documentAns.styles["Title"]
  sectPrAns = sectionAns._sectPr
  colsAns = sectPrAns.xpath('./w:cols')[0]
  colsAns.set(qn('w:num'),'3')
  count = 0
  while count != numQ:
      let = letters[randint(0,len(letters)-1)]
      #Assign random values to x, a, and b
      x = choice([i for i in range(1,20)])
      op = choice(ops)
      if op == "+":
        a = choice([i for i in range(1, 20)])
        ans = x + a
        eq = let + " + " + str(a) + " = " + str(ans)
      elif op == "-":
        a = choice([i for i in range(1, 20)])
        ans = x - a
        eq = let + " - " + str(a) + " = " + str(ans)
      elif op == "÷":
        a = choice([i for i in range(1, 15)])
        temp = x
        x = x * a
        eq = let + " ÷ " + str(temp) + " = " + str(a)
      elif op == "×":
        a = choice([i for i in range(1, 15)])
        temp = x*a
        eq = let + " × " + str(a) + " = " + str(temp) 
      if count%8 == 0 and count != 0:
          documentQ.add_paragraph()
      paraQ = documentQ.add_paragraph()
      runQ = paraQ.add_run()
      runQ.text = "Question " + str(count+1)
      runQ.underline = True
      documentQ.add_paragraph(eq)
      documentQ.add_paragraph()

      #if count%12 == 0 and count != 0:
          #documentAns.add_paragraph()
      paraAns = documentAns.add_paragraph()
      runAns = paraAns.add_run()
      runAns.text = "Question " + str(count+1)
      runAns.underline = True
      documentAns.add_paragraph(let + " = " + str(x))
      count += 1

  directory = os.getcwd()
  os.chdir(directory+'/'+"Worksheets")
    
  documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
  documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
  os.chdir(directory)




    
