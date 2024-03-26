from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
#---This is the title of the worksheet and also the name of the question and answer sheets---#
TITLE = "Solving Formulae - Unit 7 Higher"


questionsColumns = '2'

def function(numQ):
    #---Creates two documents one for the questions, one for the answers questions columns are how many columns the worksheet is split into---#
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),questionsColumns)
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE+" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')

    types = {"Circumference of Sector": ["r", "x"],
         "Area of Sector": ["r", "x"],
         "Volume of Cylinder": ["r", "h"],
         "Surface area of Cylinder": ["r", "h"],
         "Volume of Sphere": ["r"],
         "Surface area of Sphere": ["r"],
         "Area of Trapezium": ["a", "b", "h"]
         }
    count = 0
    while count != numQ:

        question = ""
        ans = 0
        #---If the questions are all the same number of lines, format it to go to the next page after n questions---#
        #if count%12 == 0 and count != 0:
            #documentQ.add_paragraph()
        #---Add the generated question and answer into the respective worksheet---#
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        documentQ.add_paragraph(question)
        documentQ.add_paragraph()

        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans))
        #---These two lines can be used to add supercript for squared or cubed numbers or units.---#
        #subscript = p.add_run("2")
        #subcript.font.superscript = True
        count += 1
        
    #---Save the two worksheets---#
    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
    
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)





    
