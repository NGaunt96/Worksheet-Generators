from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
TITLE = "Add and Subtract Decimals"

def AddSubtractDecimals(numQ):
    questionsColumns = '2'
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),questionsColumns)
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE+" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')

    
    count = 0
    while count != numQ:
        tens = [10, 100, 1000]
        op = choice(["+", "+", "+", "-"])
        a = randint(1, 1000)
        if a >= 100:
            b = randint(100, 1000)
            if op == "-":
                while b > a:
                    b = randint(100, 1000)
        elif a >= 10:
            b = randint(10, 99)
            if op == "-":
                while b > a:
                    b = randint(10, 99)
        else:
            b = randint(1,9)
            if op == "-":
                while b > a:
                    b = randint(1, 9)
    ten = choice(tens)
    a = a/ten
    b = b/ten
    question = str(a) + " " + op + " " + str(b)
    if op == "+":
        ans = round(a+b, 4)
    if op == "-":
        ans = round(a-b, 4)
    #---If the questions are all the same number of lines, format it to go to the next page after n questions---#
    if count%8 == 0 and count != 0:
        documentQ.add_paragraph()
    #---Add the generated question and answer into the respective worksheet---#
    paraQ = documentQ.add_paragraph()
    runQ = paraQ.add_run()
    runQ.text = "Question " + str(count+1)
    runQ.underline = True
    documentQ.add_paragraph(question)
    documentQ.add_paragraph()

    if count%12 == 0 and count != 0:
        documentAns.add_paragraph()
    paraAns = documentAns.add_paragraph()
    runAns = paraAns.add_run()
    runAns.text = "Question " + str(count+1)
    runAns.underline = True
    p = documentAns.add_paragraph(str(ans))
    #---These two lines can be used to add supercript for squared or cubed numbers or units.---#
    #subscript = p.add_run("2")
    #subcript.font.superscript = True
    count += 1
    
#---Save the two worksheets---#
directory = os.getcwd()
os.chdir(directory+'/'+"Worksheets")

documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
os.chdir(directory)




    
