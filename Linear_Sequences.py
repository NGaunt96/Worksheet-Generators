from random import *
from math import *
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

TITLE = "Linear Sequences"
questions = ["nth term", "first 5 terms"]


def LinearSequences(numQ):
    documentQ = Document()
    sectionQ = documentQ.sections[0]
    headerQ = sectionQ.header
    paragraphQ = headerQ.paragraphs[0]
    paragraphQ.text = TITLE
    paragraphQ.style = documentQ.styles["Title"]
    sectPrQ = sectionQ._sectPr
    colsQ = sectPrQ.xpath('./w:cols')[0]
    colsQ.set(qn('w:num'),'2')
       
    documentAns = Document()
    sectionAns = documentAns.sections[0]
    headerAns = sectionAns.header
    paragraphAns = headerAns.paragraphs[0]
    paragraphAns.text = TITLE+" Answers"
    paragraphAns.style = documentAns.styles["Title"]
    sectPrAns = sectionAns._sectPr
    colsAns = sectPrAns.xpath('./w:cols')[0]
    colsAns.set(qn('w:num'),'3')
    count = 0
    while count != numQ:
        q = choice(questions)
        coeff = choice([i for i in range(-5,5) if i not in [0]])
        const = choice([i for i in range(-10,10)])
        seq = []
        for i in range(1,6):
            seq.append(coeff*i+const)
        if const < 0:
            op = "-"
            const = const * -1
        else:
            op = "+"
        if coeff == 1:
            coeff = ""
        elif coeff == -1:
            coeff = "-"
        nth = str(coeff) + "n " + op + str(const)
        if q == "nth term":
            question1 = "Find the nth term of the following sequence: "
            question2 = ""
            for i in seq:
                question2 += str(i)
                if i != seq[len(seq)-1]:
                    question2 += ", "
            ans = nth
        else:
            question = "Find the first 5 terms in the sequence with the nth term " + nth
            ans = ""
            for i in seq:
                ans += str(i)
                if i != seq[len(seq)-1]:
                    ans += ", "
        #if count%8 == 0 and count != 0:
            #documentQ.add_paragraph()
        paraQ = documentQ.add_paragraph()
        runQ = paraQ.add_run()
        runQ.text = "Question " + str(count+1)
        runQ.underline = True
        if q == "nth term":
            documentQ.add_paragraph(question1)
            documentQ.add_paragraph(question2)
            #documentQ.add_paragraph()
        else:    
            documentQ.add_paragraph(question)
        if count%12 == 0 and count != 0:
            documentAns.add_paragraph()
        paraAns = documentAns.add_paragraph()
        runAns = paraAns.add_run()
        runAns.text = "Question " + str(count+1)
        runAns.underline = True
        p = documentAns.add_paragraph(str(ans))
        #subscript = p.add_run("2")
        #subcript.font.superscript = True
        count += 1

    directory = os.getcwd()
    os.chdir(directory+'/'+"Worksheets")
    
    documentQ.save(str(numQ) + " " + TITLE + ' Questions.docx')
    documentAns.save(str(numQ) + " " + TITLE + ' Answers.docx')
    os.chdir(directory)




    
